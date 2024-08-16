import docx
from collections import defaultdict, Counter
import os
from datetime import datetime
import re
import string
from openpyxl import Workbook

BASE_PATH = "/Users/daniel.gaimvisma.com/PycharmProjects/parsefiles/old_files"
OUTPUT_BASE_PATH = "/Users/daniel.gaimvisma.com/PycharmProjects/parsefiles/parsed_output"


def parse_docx(file_path):
    print(f"Opening document: {file_path}")
    doc = docx.Document(file_path)
    content = defaultdict(lambda: defaultdict(list))
    current_heading1 = "No Heading 1"
    current_heading2 = None

    print(f"Total paragraphs in document: {len(doc.paragraphs)}")
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: Style = {paragraph.style.name}, Text = {paragraph.text[:50]}...")
        if paragraph.style.name == 'Heading 1':
            current_heading1 = paragraph.text
            current_heading2 = None
            print(f"Found Heading 1: {current_heading1}")
        elif paragraph.style.name == 'Heading 2':
            current_heading2 = paragraph.text
            print(f"Found Heading 2: {current_heading2}")
        elif current_heading2 is not None:
            content[current_heading1][current_heading2].append(paragraph)
        elif current_heading1 != "No Heading 1":
            content[current_heading1]["_intro"].append(paragraph)

    print(f"Parsed content structure:")
    for h1, h2_content in content.items():
        print(f"  Heading 1: {h1}")
        for h2, paragraphs in h2_content.items():
            print(f"    Heading 2: {h2}")
            print(f"      Paragraphs: {len(paragraphs)}")

    return dict(content)


def create_output_folder():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    folder_name = f"parsed_files_{timestamp}"
    output_folder = os.path.join(OUTPUT_BASE_PATH, folder_name)
    os.makedirs(output_folder, exist_ok=True)
    return output_folder


def sanitize_filename(filename):
    sanitized = re.sub(r'[<>:"/\\|?*]', '', filename)
    sanitized = sanitized.replace(' ', '_')
    return sanitized[:255]


def save_parsed_content(parsed_content, output_folder, original_filename):
    base_name = os.path.splitext(original_filename)[0]
    doc_folder = os.path.join(output_folder, sanitize_filename(base_name))
    os.makedirs(doc_folder, exist_ok=True)

    for heading1, heading2_content in parsed_content.items():
        heading1_folder = os.path.join(doc_folder, sanitize_filename(heading1))
        os.makedirs(heading1_folder, exist_ok=True)

        for heading2, paragraphs in heading2_content.items():
            if heading2 == "_intro":
                file_name = "intro.docx"
            else:
                file_name = f"{sanitize_filename(heading2)}.docx"

            output_file = os.path.join(heading1_folder, file_name)

            doc = docx.Document()
            if heading2 != "_intro":
                doc.add_heading(heading2, level=2)
            for paragraph in paragraphs:
                doc.add_paragraph(paragraph.text, style=paragraph.style)

            doc.save(output_file)

    return doc_folder


def list_docx_files(directory):
    print(f"Listing .docx files in directory: {directory}")
    if not os.path.exists(directory):
        print(f"Directory does not exist: {directory}")
        return []
    docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
    print(f"Found {len(docx_files)} .docx files")
    return docx_files


def select_file(files):
    print("\nAvailable .docx files:")
    for i, file in enumerate(files, 1):
        print(f"{i}. {file}")

    while True:
        try:
            choice = int(input("\nEnter the number of the file you want to process (0 to process all files): "))
            if choice == 0:
                return files
            elif 1 <= choice <= len(files):
                return [files[choice - 1]]
            else:
                print("Invalid choice. Please try again.")
        except ValueError:
            print("Please enter a valid number.")


def count_words_in_docx(file_path):
    doc = docx.Document(file_path)
    words = []
    for para in doc.paragraphs:
        words.extend(para.text.lower().split())
    return words


def create_word_count_summary(docx_files, output_folder, min_count, max_count):
    word_count = Counter()
    for file in docx_files:
        file_path = os.path.join(BASE_PATH, file)
        words = count_words_in_docx(file_path)
        word_count.update(words)

    filtered_words = [word for word, count in word_count.items() if min_count <= count <= max_count]
    filtered_words.sort()

    wb = Workbook()
    ws = wb.active
    ws.title = "Word Count Summary"

    ws['A1'] = "Word"
    ws['B1'] = "Count"

    for row, word in enumerate(filtered_words, start=2):
        ws.cell(row=row, column=1, value=word)
        ws.cell(row=row, column=2, value=word_count[word])

    summary_file = os.path.join(output_folder, f"word_count_summary_{min_count}_to_{max_count}.xlsx")
    wb.save(summary_file)
    print(f"Word count summary saved to: {summary_file}")


def get_word_count_range():
    while True:
        try:
            min_count = int(input("Enter the minimum word count (default is 5): ") or "5")
            max_count = int(input("Enter the maximum word count (default is 80): ") or "80")
            if min_count > 0 and max_count >= min_count:
                return min_count, max_count
            else:
                print(
                    "Invalid range. Minimum count should be positive and maximum count should be greater than or equal to minimum count.")
        except ValueError:
            print("Please enter valid numbers.")


def select_features():
    features = []
    print("\nSelect features to run:")
    print("1. Parse documents")
    print("2. Create word count summary")
    while True:
        choice = input("Enter the numbers of the features you want to run (e.g., '1 2' for both): ")
        selected = [int(c) for c in choice.split() if c.isdigit() and 1 <= int(c) <= 2]
        if selected:
            return selected
        print("Invalid selection. Please try again.")


def main():
    print(f"Documents will be read from: {BASE_PATH}")
    print(f"Current working directory: {os.getcwd()}")
    print(f"BASE_PATH exists: {os.path.exists(BASE_PATH)}")

    all_docx_files = list_docx_files(BASE_PATH)

    if not all_docx_files:
        print(f"No .docx files found in {BASE_PATH}")
        return

    selected_files = select_file(all_docx_files)
    selected_features = select_features()
    output_folder = create_output_folder()

    try:
        if 1 in selected_features:  # Parse documents
            for file in selected_files:
                file_path = os.path.join(BASE_PATH, file)
                parsed_content = parse_docx(file_path)
                if not parsed_content:
                    print(f"No content found in the document: {file}")
                else:
                    saved_folder = save_parsed_content(parsed_content, output_folder, file)
                    print(f"\nParsed content for {file} has been saved to: {saved_folder}")
                    print("\nContent overview:")
                    for heading1, heading2_content in parsed_content.items():
                        print(f"\nHeading 1: {heading1}")
                        for heading2, paragraphs in heading2_content.items():
                            if heading2 == "_intro":
                                print(f"  Introduction: {len(paragraphs)} paragraphs")
                            else:
                                print(f"  Heading 2: {heading2}")
                                print(f"    {len(paragraphs)} paragraphs")

        if 2 in selected_features:  # Create word count summary
            min_count, max_count = get_word_count_range()
            create_word_count_summary(selected_files, output_folder, min_count, max_count)

    except Exception as e:
        print(f"An error occurred while processing the file(s): {e}")
        print(f"Error type: {type(e).__name__}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()