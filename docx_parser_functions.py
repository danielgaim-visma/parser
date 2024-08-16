import docx
from collections import OrderedDict, Counter
import os
import re
import logging
from openpyxl import Workbook
import csv

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


def read_keywords(reference_file):
    keywords = []
    file_extension = os.path.splitext(reference_file)[1].lower()

    try:
        if file_extension == '.csv':
            with open(reference_file, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                keywords = [row[0].strip().lower() for row in reader if row]
        elif file_extension == '.txt':
            with open(reference_file, 'r', encoding='utf-8') as file:
                keywords = [line.strip().lower() for line in file if line.strip()]
        else:
            logger.error(f"Unsupported file format for reference file: {file_extension}")
            return []
    except Exception as e:
        logger.error(f"Error reading reference file: {str(e)}")
        return []

    return keywords


def tag_document(doc_path, keywords):
    try:
        doc = docx.Document(doc_path)
        text = " ".join([para.text.lower() for para in doc.paragraphs])

        tags = [keyword for keyword in keywords if keyword in text]

        if tags:
            # Add a new paragraph for tags
            doc.add_paragraph()  # Add an empty paragraph for spacing
            tag_paragraph = doc.add_paragraph("Tags:")
            for tag in tags:
                tag_paragraph.add_run(f" tag {tag}")

            doc.save(doc_path)

        return tags
    except Exception as e:
        logger.error(f"Error tagging document {doc_path}: {str(e)}")
        return []


def parse_docx(file_path, output_folder, parse_level, keywords):
    logger.info(f"Parsing document: {file_path}")
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        logger.error(f"Error opening document {file_path}: {str(e)}")
        raise

    content = OrderedDict()
    current_headings = [None, None, None]
    current_content = []

    # Create main folder for the document
    doc_title = os.path.splitext(os.path.basename(file_path))[0]
    doc_folder = os.path.join(output_folder, sanitize_filename(doc_title))
    os.makedirs(doc_folder, exist_ok=True)

    def save_current_content():
        if current_headings[0]:
            if current_headings[0] not in content:
                content[current_headings[0]] = OrderedDict()
            if current_headings[1]:
                if current_headings[1] not in content[current_headings[0]]:
                    content[current_headings[0]][current_headings[1]] = OrderedDict()
                if current_headings[2]:
                    content[current_headings[0]][current_headings[1]][current_headings[2]] = current_content.copy()
                else:
                    content[current_headings[0]][current_headings[1]][""] = current_content.copy()
            else:
                content[current_headings[0]][""] = {"": current_content.copy()}
        current_content.clear()

    for paragraph in doc.paragraphs:
        heading_level = get_heading_level(paragraph)

        if heading_level and heading_level <= 3:
            save_current_content()
            current_headings[heading_level - 1] = paragraph.text
            for i in range(heading_level, 3):
                current_headings[i] = None
        else:
            current_content.append(paragraph)

    save_current_content()

    # Write content to DOCX files
    for h1, h2_dict in content.items():
        h1_folder = os.path.join(doc_folder, sanitize_filename(h1))
        os.makedirs(h1_folder, exist_ok=True)

        if parse_level == 1:
            file_name = f"{sanitize_filename(h1)}.docx"
            full_path = os.path.join(doc_folder, file_name)
            save_docx(full_path, h1, "", "", h2_dict[""][""], 1)
            tags = tag_document(full_path, keywords)
            logger.info(f"Tagged document {full_path} with tags: {tags}")
        else:
            for h2, h3_dict in h2_dict.items():
                if h2:
                    h2_folder = os.path.join(h1_folder, sanitize_filename(h2))
                    os.makedirs(h2_folder, exist_ok=True)

                    if parse_level == 2:
                        file_name = f"{sanitize_filename(h2)}.docx"
                        full_path = os.path.join(h1_folder, file_name)
                        save_docx(full_path, h1, h2, "", h3_dict[""], 2)
                        tags = tag_document(full_path, keywords)
                        logger.info(f"Tagged document {full_path} with tags: {tags}")
                    elif parse_level == 3:
                        for h3, paragraphs in h3_dict.items():
                            if h3:
                                file_name = f"{sanitize_filename(h3)}.docx"
                                full_path = os.path.join(h2_folder, file_name)
                                save_docx(full_path, h1, h2, h3, paragraphs, 3)
                                tags = tag_document(full_path, keywords)
                                logger.info(f"Tagged document {full_path} with tags: {tags}")
                            else:
                                file_name = f"{sanitize_filename(h2)}_content.docx"
                                full_path = os.path.join(h2_folder, file_name)
                                save_docx(full_path, h1, h2, "", paragraphs, 2)
                                tags = tag_document(full_path, keywords)
                                logger.info(f"Tagged document {full_path} with tags: {tags}")

    logger.info(f"Parsing complete. Output folder: {doc_folder}")
    return doc_folder


def save_docx(full_path, h1, h2, h3, paragraphs, level):
    try:
        doc = docx.Document()
        if level >= 1 and h1:
            doc.add_heading(h1, level=1)
        if level >= 2 and h2:
            doc.add_heading(h2, level=2)
        if level == 3 and h3:
            doc.add_heading(h3, level=3)

        for para in paragraphs:
            new_para = doc.add_paragraph()
            new_para.style = para.style
            new_para.text = para.text
            # Copy runs to preserve formatting
            new_para.runs.clear()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                # Add more formatting attributes as needed

        doc.save(full_path)
        logger.info(f"Saved parsed content to: {full_path}")
    except Exception as e:
        logger.error(f"Error saving document {full_path}: {str(e)}")


def get_heading_level(paragraph):
    if paragraph.style.name.startswith('Heading'):
        try:
            return int(paragraph.style.name.split()[-1])
        except ValueError:
            logger.warning(f"Invalid heading style: {paragraph.style.name}")
    return None


def sanitize_filename(filename):
    # Remove invalid characters and limit length
    sanitized = re.sub(r'[^\w\-_\. ]', '_', filename)
    return sanitized[:255]  # Limit to 255 characters


# ... (create_word_count_summary function remains unchanged)


def create_word_count_summary(doc_paths, output_folder, min_count=20, max_count=100):
    logger.info(
        f"Creating word count summary for {len(doc_paths)} documents with min_count={min_count}, max_count={max_count}")
    word_count = Counter()

    for doc_path in doc_paths:
        try:
            doc = docx.Document(doc_path)
            for para in doc.paragraphs:
                words = para.text.lower().split()
                word_count.update(words)
        except Exception as e:
            logger.error(f"Error processing file {doc_path} for word count: {str(e)}")

    filtered_words = {word: count for word, count in word_count.items() if min_count <= count <= max_count}

    if not filtered_words:
        min_count_found = min(word_count.values()) if word_count else 0
        max_count_found = max(word_count.values()) if word_count else 0
        logger.warning(
            f"No words match the specified count range. Word count range in documents: {min_count_found} - {max_count_found}")
        return None, f"No words found with count between {min_count} and {max_count}. Word count range in documents: {min_count_found} - {max_count_found}"

    summary_filename = f"word_count_summary_{min_count}_to_{max_count}.xlsx"
    summary_path = os.path.join(output_folder, summary_filename)

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Word Count Summary"
        ws.append(["Word", "Count"])
        for word, count in sorted(filtered_words.items(), key=lambda x: x[1], reverse=True):
            ws.append([word, count])
        wb.save(summary_path)
        logger.info(f"Word count summary saved to: {summary_path}")
    except Exception as e:
        logger.error(f"Error writing word count summary: {str(e)}")
        return None, f"Error creating summary: {str(e)}"

    return summary_path, f"Word count summary created with {len(filtered_words)} words"