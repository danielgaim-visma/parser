import os
import re
from docx import Document
from collections import OrderedDict
import logging
from .keyword_tagger import tag_document
from concurrent.futures import ThreadPoolExecutor, as_completed

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def parse_multiple_docx(file_paths, output_folder, parse_level, keywords):
    results = []
    with ThreadPoolExecutor() as executor:
        future_to_file = {executor.submit(parse_docx, file_path, output_folder, parse_level, keywords): file_path for file_path in file_paths}
        for future in as_completed(future_to_file):
            file_path = future_to_file[future]
            try:
                doc_folder = future.result()
                results.append({"file": os.path.basename(file_path), "output_folder": doc_folder})
            except Exception as exc:
                logger.error(f'{file_path} generated an exception: {exc}')
                results.append({"file": os.path.basename(file_path), "error": str(exc)})
    return results

def parse_docx(file_path, output_folder, parse_level, keywords):
    logger.info(f"Parsing document: {file_path}")
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"Error opening document {file_path}: {str(e)}")
        raise

    content = OrderedDict()
    current_headings = [None, None, None]
    current_content = []

    def save_current_content():
        if current_headings[0]:
            if current_headings[0] not in content:
                content[current_headings[0]] = OrderedDict()
            if current_headings[1]:
                if current_headings[1] not in content[current_headings[0]]:
                    content[current_headings[0]][current_headings[1]] = OrderedDict()
                if current_headings[2]:
                    content[current_headings[0]][current_headings[1]][current_headings[2]] = current_content.copy()
                else:
                    content[current_headings[0]][current_headings[1]][""] = current_content.copy()
            else:
                content[current_headings[0]][""] = {"": current_content.copy()}
        current_content.clear()

    for paragraph in doc.paragraphs:
        heading_level = get_heading_level(paragraph)

        if heading_level and heading_level <= 3:
            save_current_content()
            current_headings[heading_level - 1] = paragraph.text
            for i in range(heading_level, 3):
                current_headings[i] = None
        else:
            current_content.append(paragraph)

    save_current_content()

    # Create main folder for the document
    doc_title = os.path.splitext(os.path.basename(file_path))[0]
    doc_folder = os.path.join(output_folder, sanitize_filename(doc_title))
    os.makedirs(doc_folder, exist_ok=True)

    # Write content to DOCX files
    for h1, h2_dict in content.items():
        h1_folder = os.path.join(doc_folder, sanitize_filename(h1))
        os.makedirs(h1_folder, exist_ok=True)

        if parse_level == 1:
            file_name = f"{sanitize_filename(h1)}.docx"
            full_path = os.path.join(doc_folder, file_name)
            save_docx(full_path, h1, "", "", h2_dict, 1)
            if keywords:
                tags = tag_document(full_path, keywords)
                logger.info(f"Tagged document {full_path} with tags: {tags}")
        else:
            for h2, h3_dict in h2_dict.items():
                h2_folder = os.path.join(h1_folder, sanitize_filename(h2)) if h2 else h1_folder
                os.makedirs(h2_folder, exist_ok=True)

                if parse_level == 2:
                    file_name = f"{sanitize_filename(h2 or 'content')}.docx"
                    full_path = os.path.join(h2_folder, file_name)
                    save_docx(full_path, h1, h2, "", h3_dict, 2)
                    if keywords:
                        tags = tag_document(full_path, keywords)
                        logger.info(f"Tagged document {full_path} with tags: {tags}")
                elif parse_level == 3:
                    for h3, paragraphs in h3_dict.items():
                        file_name = f"{sanitize_filename(h3 or 'content')}.docx"
                        full_path = os.path.join(h2_folder, file_name)
                        save_docx(full_path, h1, h2, h3, paragraphs, 3)
                        if keywords:
                            tags = tag_document(full_path, keywords)
                            logger.info(f"Tagged document {full_path} with tags: {tags}")

    logger.info(f"Parsing complete. Output folder: {doc_folder}")
    return doc_folder

def save_docx(full_path, h1, h2, h3, content, level):
    try:
        doc = Document()
        if level >= 1 and h1:
            doc.add_heading(h1, level=1)
        if level >= 2 and h2:
            doc.add_heading(h2, level=2)
        if level == 3 and h3:
            doc.add_heading(h3, level=3)

        if level == 1:
            for h2, h3_dict in content.items():
                if h2:
                    doc.add_heading(h2, level=2)
                for h3, paragraphs in h3_dict.items():
                    if h3:
                        doc.add_heading(h3, level=3)
                    add_paragraphs(doc, paragraphs)
        elif level == 2:
            for h3, paragraphs in content.items():
                if h3:
                    doc.add_heading(h3, level=3)
                add_paragraphs(doc, paragraphs)
        else:
            add_paragraphs(doc, content)

        doc.save(full_path)
        logger.info(f"Saved parsed content to: {full_path}")
    except Exception as e:
        logger.error(f"Error saving document {full_path}: {str(e)}")

def add_paragraphs(doc, paragraphs):
    for para in paragraphs:
        new_para = doc.add_paragraph()
        new_para.style = para.style
        new_para.text = para.text
        # Copy runs to preserve formatting
        new_para.runs.clear()
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            # Add more formatting attributes as needed

def get_heading_level(paragraph):
    if paragraph.style.name.startswith('Heading'):
        try:
            return int(paragraph.style.name.split()[-1])
        except ValueError:
            logger.warning(f"Invalid heading style: {paragraph.style.name}")
    return None

def sanitize_filename(filename):
    # Remove invalid characters and limit length
    sanitized = re.sub(r'[^\w\-_\. ]', '_', filename)
    return sanitized[:255]  # Limit to 255 characters