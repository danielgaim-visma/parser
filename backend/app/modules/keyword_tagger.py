import os
import csv
from docx import Document
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed

logger = logging.getLogger(__name__)

def read_keywords(reference_file):
    keywords = []
    file_extension = os.path.splitext(reference_file)[1].lower()

    try:
        if file_extension == '.csv':
            with open(reference_file, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                keywords = [row[0].strip().lower() for row in reader if row]
        elif file_extension == '.txt':
            with open(reference_file, 'r', encoding='utf-8') as file:
                keywords = [line.strip().lower() for line in file if line.strip()]
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        logger.info(f"Successfully read {len(keywords)} keywords from {reference_file}")
    except Exception as e:
        logger.error(f"Error reading reference file: {str(e)}")
        raise

    return keywords

def tag_document(doc_path, keywords):
    try:
        doc = Document(doc_path)
        text = " ".join([para.text.lower() for para in doc.paragraphs])

        tags = [keyword for keyword in keywords if keyword in text]

        if tags:
            doc.add_paragraph()
            tag_paragraph = doc.add_paragraph("Tags: ")
            tag_paragraph.add_run(', '.join(f'"{tag}"' for tag in tags))

            doc.save(doc_path)
            logger.info(f"Document {doc_path} tagged with {len(tags)} keywords")
        else:
            logger.info(f"No matching keywords found in document {doc_path}")

        return tags
    except Exception as e:
        logger.error(f"Error tagging document {doc_path}: {str(e)}")
        raise

def tag_multiple_documents(doc_paths, output_folder, keywords):
    results = []
    with ThreadPoolExecutor() as executor:
        future_to_file = {executor.submit(tag_document_and_save, doc_path, output_folder, keywords): doc_path for doc_path in doc_paths}
        for future in as_completed(future_to_file):
            doc_path = future_to_file[future]
            try:
                output_file, tags = future.result()
                results.append({"file": os.path.basename(doc_path), "output_file": output_file, "tags": tags})
                logger.info(f"Successfully processed {doc_path}")
            except Exception as exc:
                logger.error(f'Error processing {doc_path}: {exc}')
                results.append({"file": os.path.basename(doc_path), "error": str(exc)})
    return results

def tag_document_and_save(doc_path, output_folder, keywords):
    try:
        doc = Document(doc_path)
        text = " ".join([para.text.lower() for para in doc.paragraphs])

        tags = [keyword for keyword in keywords if keyword in text]

        if tags:
            doc.add_paragraph()
            tag_paragraph = doc.add_paragraph("Tags: ")
            tag_paragraph.add_run(', '.join(f'"{tag}"' for tag in tags))

        output_filename = f"tagged_{os.path.basename(doc_path)}"
        output_path = os.path.join(output_folder, output_filename)
        doc.save(output_path)

        logger.info(f"Tagged document saved: {output_path}")
        return output_filename, tags
    except Exception as e:
        logger.error(f"Error tagging and saving document {doc_path}: {str(e)}")
        raise